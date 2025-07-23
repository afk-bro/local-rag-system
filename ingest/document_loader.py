"""
Document loaders for various file formats.
Supports PDF, DOCX, HTML, TXT, MD, JSON, CSV, and YouTube transcripts.
"""

from typing import List, Dict, Any, Optional
from pathlib import Path
import logging
from langchain.schema import Document
from langchain.document_loaders import (
    PyPDFLoader,
    TextLoader,
    UnstructuredHTMLLoader
)
import pandas as pd
import json
from docx import Document as DocxDocument
from bs4 import BeautifulSoup
from youtube_transcript_api import YouTubeTranscriptApi
import re
from config.settings import config

logger = logging.getLogger(__name__)


class DocumentLoader:
    """
    Universal document loader that handles multiple file formats.
    """
    
    def __init__(self):
        """Initialize the document loader."""
        self.supported_formats = config.supported_formats
        
    def load_document(self, file_path: str) -> List[Document]:
        """
        Load a document from file path.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            List of Document objects
            
        Raises:
            ValueError: If file format is not supported
            FileNotFoundError: If file doesn't exist
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_extension = path.suffix.lower().lstrip('.')
        
        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        # Route to appropriate loader based on file extension
        loader_map = {
            'pdf': self._load_pdf,
            'txt': self._load_text,
            'md': self._load_text,
            'docx': self._load_docx,
            'html': self._load_html,
            'json': self._load_json,
            'csv': self._load_csv
        }
        
        loader_func = loader_map.get(file_extension)
        if not loader_func:
            raise ValueError(f"No loader implemented for: {file_extension}")
        
        try:
            documents = loader_func(str(path))
            logger.info(f"Loaded {len(documents)} documents from {file_path}")
            return documents
        except Exception as e:
            logger.error(f"Error loading {file_path}: {str(e)}")
            raise
    
    def _load_pdf(self, file_path: str) -> List[Document]:
        """Load PDF document."""
        loader = PyPDFLoader(file_path)
        documents = loader.load()
        
        # Add metadata
        for doc in documents:
            doc.metadata.update({
                'source': file_path,
                'file_type': 'pdf',
                'loader': 'PyPDFLoader'
            })
        
        return documents
    
    def _load_text(self, file_path: str) -> List[Document]:
        """Load text/markdown document."""
        loader = TextLoader(file_path, encoding='utf-8')
        documents = loader.load()
        
        # Determine if it's markdown
        file_type = 'markdown' if file_path.endswith('.md') else 'text'
        
        for doc in documents:
            doc.metadata.update({
                'source': file_path,
                'file_type': file_type,
                'loader': 'TextLoader'
            })
        
        return documents
    
    def _load_docx(self, file_path: str) -> List[Document]:
        """Load DOCX document."""
        doc = DocxDocument(file_path)
        
        # Extract text from paragraphs
        paragraphs = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                paragraphs.append(paragraph.text)
        
        content = '\n\n'.join(paragraphs)
        
        document = Document(
            page_content=content,
            metadata={
                'source': file_path,
                'file_type': 'docx',
                'loader': 'python-docx',
                'paragraph_count': len(paragraphs)
            }
        )
        
        return [document]
    
    def _load_html(self, file_path: str) -> List[Document]:
        """Load HTML document."""
        with open(file_path, 'r', encoding='utf-8') as file:
            html_content = file.read()
        
        # Parse HTML and extract text
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
        
        # Extract text
        text = soup.get_text()
        
        # Clean up whitespace
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = ' '.join(chunk for chunk in chunks if chunk)
        
        document = Document(
            page_content=text,
            metadata={
                'source': file_path,
                'file_type': 'html',
                'loader': 'BeautifulSoup',
                'title': soup.title.string if soup.title else None
            }
        )
        
        return [document]
    
    def _load_json(self, file_path: str) -> List[Document]:
        """Load JSON document."""
        with open(file_path, 'r', encoding='utf-8') as file:
            data = json.load(file)
        
        # Handle different JSON structures
        if isinstance(data, list):
            documents = []
            for i, item in enumerate(data):
                content = json.dumps(item, indent=2) if isinstance(item, dict) else str(item)
                doc = Document(
                    page_content=content,
                    metadata={
                        'source': file_path,
                        'file_type': 'json',
                        'loader': 'json',
                        'item_index': i,
                        'total_items': len(data)
                    }
                )
                documents.append(doc)
            return documents
        else:
            content = json.dumps(data, indent=2)
            document = Document(
                page_content=content,
                metadata={
                    'source': file_path,
                    'file_type': 'json',
                    'loader': 'json'
                }
            )
            return [document]
    
    def _load_csv(self, file_path: str) -> List[Document]:
        """Load CSV document."""
        df = pd.read_csv(file_path)
        
        # Convert DataFrame to text representation
        content = df.to_string(index=False)
        
        document = Document(
            page_content=content,
            metadata={
                'source': file_path,
                'file_type': 'csv',
                'loader': 'pandas',
                'rows': len(df),
                'columns': len(df.columns),
                'column_names': list(df.columns)
            }
        )
        
        return [document]
    
    def load_youtube_transcript(self, video_url: str, language: str = 'en') -> List[Document]:
        """
        Load YouTube video transcript.
        
        Args:
            video_url: YouTube video URL
            language: Language code for transcript
            
        Returns:
            List of Document objects
        """
        # Extract video ID from URL
        video_id = self._extract_video_id(video_url)
        if not video_id:
            raise ValueError(f"Could not extract video ID from URL: {video_url}")
        
        try:
            # Get transcript
            transcript = YouTubeTranscriptApi.get_transcript(video_id, languages=[language])
            
            # Combine transcript segments
            full_text = ' '.join([entry['text'] for entry in transcript])
            
            document = Document(
                page_content=full_text,
                metadata={
                    'source': video_url,
                    'video_id': video_id,
                    'file_type': 'youtube_transcript',
                    'loader': 'youtube-transcript-api',
                    'language': language,
                    'duration': transcript[-1]['start'] + transcript[-1]['duration'] if transcript else 0,
                    'segment_count': len(transcript)
                }
            )
            
            return [document]
            
        except Exception as e:
            logger.error(f"Error loading YouTube transcript for {video_url}: {str(e)}")
            raise
    
    def _extract_video_id(self, url: str) -> Optional[str]:
        """Extract YouTube video ID from URL."""
        patterns = [
            r'(?:youtube\.com/watch\?v=|youtu\.be/|youtube\.com/embed/)([^&\n?#]+)',
            r'youtube\.com/watch\?.*v=([^&\n?#]+)'
        ]
        
        for pattern in patterns:
            match = re.search(pattern, url)
            if match:
                return match.group(1)
        
        return None
    
    def load_directory(self, directory_path: str, recursive: bool = True) -> List[Document]:
        """
        Load all supported documents from a directory.
        
        Args:
            directory_path: Path to directory
            recursive: Whether to search recursively
            
        Returns:
            List of all loaded documents
        """
        path = Path(directory_path)
        if not path.exists():
            raise FileNotFoundError(f"Directory not found: {directory_path}")
        
        documents = []
        pattern = "**/*" if recursive else "*"
        
        for file_path in path.glob(pattern):
            if file_path.is_file():
                file_extension = file_path.suffix.lower().lstrip('.')
                if file_extension in self.supported_formats:
                    try:
                        docs = self.load_document(str(file_path))
                        documents.extend(docs)
                    except Exception as e:
                        logger.warning(f"Failed to load {file_path}: {str(e)}")
                        continue
        
        logger.info(f"Loaded {len(documents)} documents from {directory_path}")
        return documents


def create_document_loader() -> DocumentLoader:
    """Factory function to create a document loader."""
    return DocumentLoader()
